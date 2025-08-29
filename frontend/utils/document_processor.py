import PyPDF2
from docx import Document
import os
import re
import tempfile
import json
from PIL import Image
import pytesseract
import pdfplumber
import pandas as pd

from models.get_llm import get_llm


class DocumentProcessor:
    def __init__(self, docs_folder):
        self.docs_folder = docs_folder
        self.rules = []
        self.templates = []
        # 缓存LLM实例
        self._llm = None
        # 缓存已处理的内容
        self._summaries_cache = {}
        self._structures_cache = {}
        # 新增：记录处理的文件列表与错误信息（用于Streamlit反馈）
        self.processed_rule_files = []  # 处理成功的规则文档名
        self.processed_template_files = []  # 处理成功的模板文档名
        self.process_errors = []  # 解析错误的文档信息（格式："文件名：错误原因"）

        # -------------------------- 新增：图片OCR配置（需用户手动确认路径）--------------------------
        # 1. Windows用户：需将路径改为你的Tesseract安装路径（默认如下，若修改过安装位置需调整）
        self.tesseract_path = r'D:\Program Files\Tesseract-OCR\tesseract.exe'
        # 2. Mac用户：注释上面一行，启用下面一行（默认路径，若用brew安装）
        # self.tesseract_path = '/usr/local/bin/tesseract'

        # 配置Tesseract路径
        try:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_path
            # 验证Tesseract是否可用（避免后续OCR时才报错）
            pytesseract.get_tesseract_version()
        except Exception as e:
            self.process_errors.append(
                f"Tesseract配置错误：{str(e)}，图片OCR功能将不可用。请检查安装路径或重新安装Tesseract（附安装指南：https://github.com/UB-Mannheim/tesseract/wiki）")

    def _get_llm_instance(self):
        """获取LLM实例（单例模式）——原有方法，无修改"""
        if self._llm is None:
            self._llm = get_llm()
        return self._llm

    # -------------------------- 新增：核心工具方法（图片OCR+表格转换）--------------------------
    def _ocr_single_image(self, image_path):
        """单张图片OCR识别，返回结构化文本（标注图片来源）"""
        try:
            # 图片预处理：转为灰度图提升识别率（减少彩色干扰）
            with Image.open(image_path) as img:
                img_gray = img.convert('L')
                # 执行OCR（支持中英双语，避免金融术语识别错误）
                ocr_text = pytesseract.image_to_string(
                    img_gray,
                    lang='chi_sim+eng',  # 中文+英文识别（必须安装对应语言包）
                    config='--psm 6'  # 按段落识别，适合文档类图片
                )
            # 结构化输出：标注图片OCR结果，方便后续LLM识别
            return f"【图片OCR内容】\n{ocr_text.strip()}\n【图片OCR结束】" if ocr_text.strip() else "【图片OCR内容】未识别到有效文字\n【图片OCR结束】"
        except Exception as e:
            return f"【图片OCR内容】识别失败：{str(e)}\n【图片OCR结束】"

    def _table_to_markdown(self, table_data, table_idx, file_type):
        """将表格数据（列表/元组）转为Markdown格式（LLM易识别结构）"""
        try:
            # 处理空表格
            if not table_data or len(table_data) < 1:
                return f"【{file_type}表格{table_idx}】空表格，无数据"

            # 用Pandas整理表格（自动处理空值、对齐列数）
            df = pd.DataFrame(table_data)
            # 若第一行是表头，设为列名（避免表头和内容混为一谈）
            if len(df.columns) > 1 and all(isinstance(cell, str) and len(cell.strip()) > 0 for cell in df.iloc[0]):
                df.columns = df.iloc[0]
                df = df.drop(0).reset_index(drop=True)

            # 空值替换为空白字符串（避免显示NaN）
            df = df.fillna("")
            # 转为Markdown表格
            markdown_table = df.to_markdown(index=False, tablefmt="pipe")
            return f"【{file_type}表格{table_idx}】\n{markdown_table}\n"
        except Exception as e:
            return f"【{file_type}表格{table_idx}】转换失败：{str(e)}\n原始表格数据：{str(table_data[:3])}..."  # 只显示前3行避免过长

    # -------------------------- 新增：PDF图片+表格处理（替换原有纯文本提取）--------------------------
    def _process_pdf_with_table_image(self, file_path):
        """处理PDF中的文本+表格+图片，返回整合后结构化内容"""
        file_name = os.path.basename(file_path)
        content_parts = [f"【PDF文件：{file_name} 内容开始】"]

        try:
            # 1. 提取PDF纯文本（保留基础内容，用于补充表格/图片外的信息）
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_content = ""
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_content += f"【第{page_num}页文本】\n{page_text.strip()}\n"
            if text_content:
                content_parts.append(text_content)

            # 2. 用pdfplumber提取PDF表格（比PyPDF2更精准，保留行列结构）
            with pdfplumber.open(file_path) as pdf:
                table_count = 0
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables()  # 提取当前页所有表格
                    if tables:
                        for table in tables:
                            table_count += 1
                            # 表格转Markdown并添加到内容
                            table_markdown = self._table_to_markdown(table, table_count, "PDF")
                            content_parts.append(f"【第{page_num}页表格】\n{table_markdown}")
                if table_count == 0:
                    content_parts.append("【PDF表格信息】未检测到表格\n")

            # 3. 提取PDF中的图片并执行OCR
            with pdfplumber.open(file_path) as pdf:
                img_count = 0
                for page_num, page in enumerate(pdf.pages, 1):
                    images = page.images  # 提取当前页所有图片
                    if images:
                        for img_info in images:
                            img_count += 1
                            # 提取图片二进制数据，保存为临时文件（OCR完成后删除）
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                                temp_img.write(img_info['stream'].read())
                                temp_img_path = temp_img.name

                            # 执行OCR并添加到内容
                            ocr_result = self._ocr_single_image(temp_img_path)
                            content_parts.append(f"【第{page_num}页图片{img_count}】\n{ocr_result}\n")

                            # 删除临时图片（避免占用磁盘空间）
                            os.unlink(temp_img_path)
                if img_count == 0:
                    content_parts.append("【PDF图片信息】未检测到图片\n")

            # 整合所有内容，返回结构化文本
            content_parts.append(f"【PDF文件：{file_name} 内容结束】")
            return "\n".join(content_parts)

        except Exception as e:
            # 若处理失败，降级为原有纯文本提取（避免完全无法使用）
            err_msg = f"PDF高级处理（表格+图片）失败：{str(e)}，已降级为纯文本提取"
            self.process_errors.append(f"{file_name}：{err_msg}")
            # 调用原有纯文本提取逻辑
            pure_text = self._process_pdf_fallback(file_path)
            return f"【PDF文件：{file_name} 内容（降级纯文本）】\n{pure_text}\n【处理说明】{err_msg}"

    def _process_pdf_fallback(self, file_path):
        """PDF高级处理失败时的降级方案——原有纯文本提取逻辑，无修改"""
        content = ""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted_text = page.extract_text() or ""
                    content += extracted_text
        except Exception as e:
            err_msg = f"纯文本提取也失败：{str(e)}"
            self.process_errors.append(f"{os.path.basename(file_path)}：{err_msg}")
            content = err_msg
        return content

    # -------------------------- 新增：Word图片+表格处理（替换原有纯文本提取）--------------------------
    def _process_docx_with_table_image(self, file_path):
        """处理Word中的文本+表格+图片，返回整合后结构化内容"""
        file_name = os.path.basename(file_path)
        content_parts = [f"【Word文件：{file_name} 内容开始】"]

        try:
            doc = Document(file_path)

            # 1. 提取Word段落纯文本
            text_content = ""
            for para_idx, para in enumerate(doc.paragraphs, 1):
                para_text = para.text.strip()
                if para_text:
                    text_content += f"【段落{para_idx}】\n{para_text}\n"
            if text_content:
                content_parts.append(text_content)
            else:
                content_parts.append("【Word文本信息】未检测到有效段落文本\n")

            # 2. 提取Word表格
            if doc.tables:
                for table_idx, table in enumerate(doc.tables, 1):
                    # 读取表格数据（行→列）
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]  # 清理换行符
                        table_data.append(row_data)
                    # 表格转Markdown
                    table_markdown = self._table_to_markdown(table_data, table_idx, "Word")
                    content_parts.append(f"【表格{table_idx}】\n{table_markdown}\n")
            else:
                content_parts.append("【Word表格信息】未检测到表格\n")

            # 3. 提取Word中的图片（需通过文档关系获取）
            img_count = 0
            # 遍历文档中的所有关系（图片存储在关系中）
            for rel_id, rel in doc.part.related_parts.items():
                # 判断是否为图片类型（MIME类型匹配）
                if rel.content_type.startswith('image/'):
                    img_count += 1
                    # 保存图片到临时文件
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                        temp_img.write(rel.blob)  # 写入图片二进制数据
                        temp_img_path = temp_img.name

                    # 执行OCR
                    ocr_result = self._ocr_single_image(temp_img_path)
                    content_parts.append(f"【图片{img_count}】\n{ocr_result}\n")

                    # 删除临时图片
                    os.unlink(temp_img_path)
            if img_count == 0:
                content_parts.append("【Word图片信息】未检测到图片\n")

            # 整合所有内容
            content_parts.append(f"【Word文件：{file_name} 内容结束】")
            return "\n".join(content_parts)

        except Exception as e:
            # 降级为原有纯文本提取
            err_msg = f"Word高级处理（表格+图片）失败：{str(e)}，已降级为纯文本提取"
            self.process_errors.append(f"{file_name}：{err_msg}")
            pure_text = self._process_docx_fallback(file_path)
            return f"【Word文件：{file_name} 内容（降级纯文本）】\n{pure_text}\n【处理说明】{err_msg}"

    def _process_docx_fallback(self, file_path):
        """Word高级处理失败时的降级方案——原有纯文本提取逻辑，无修改"""
        content = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        except Exception as e:
            err_msg = f"纯文本提取也失败：{str(e)}"
            self.process_errors.append(f"{os.path.basename(file_path)}：{err_msg}")
            content = err_msg
        return content

    # -------------------------- 原有方法：仅修改内部调用（接口不变，其他代码无需改）--------------------------
    def _process_pdf(self, file_path):
        """原有PDF处理方法——仅修改内部调用，接口（输入输出）完全不变"""
        return self._process_pdf_with_table_image(file_path)  # 调用新增的高级处理方法

    def _process_docx(self, file_path):
        """原有Word处理方法——仅修改内部调用，接口（输入输出）完全不变"""
        return self._process_docx_with_table_image(file_path)  # 调用新增的高级处理方法

    def _is_rule_document(self, file_name, content):
        """判断文档是否为规则文档——原有方法，无修改"""
        rule_keywords = ['规则', '规范', '指引', '要求', '监管', '条例', 'regulation', 'rule', 'guideline',
                         'requirement']
        file_name_lower = file_name.lower()
        content_lower = content.lower() if content else ""

        file_keyword_count = sum(1 for kw in rule_keywords if kw.lower() in file_name_lower)
        if file_keyword_count >= 2:
            return True
        elif file_keyword_count == 1:
            content_keyword_count = sum(1 for kw in rule_keywords if kw.lower() in content_lower)
            return content_keyword_count >= 1
        else:
            content_keyword_count = sum(1 for kw in rule_keywords if kw.lower() in content_lower)
            return content_keyword_count >= 2

    def _summarize_content(self, content):
        """使用LLM生成文档摘要——原有方法，无修改（LLM会自动识别新增的表格/图片标注）"""
        cache_key = hash(content)
        if cache_key in self._summaries_cache:
            return self._summaries_cache[cache_key]

        try:
            llm = self._get_llm_instance()
            if len(content) <= 2000:
                prompt = [
                    {"role": "user",
                     "content": f"请总结以下金融SOA相关文档的核心内容，重点提取：1.规则要求 2.模板结构 3.必填模块 4.表格中的关键数据（如费用标准、风险等级）5.图片OCR中的有效信息，摘要需简洁完整：{content}"}
                ]
                result = llm(prompt)
            elif 2001 <= len(content) <= 10000:
                import re
                chapter_pattern = re.compile(r'(\d+\.\s+|一、\s+|二、\s+|###\s+)')
                chunks = chapter_pattern.split(content)
                full_chunks = []
                current_chunk = ""
                for part in chunks:
                    if chapter_pattern.match(part):
                        if current_chunk:
                            full_chunks.append(current_chunk)
                        current_chunk = part
                    else:
                        current_chunk += part
                if current_chunk:
                    full_chunks.append(current_chunk)
                summaries = []
                for idx, chunk in enumerate(full_chunks, 1):
                    if len(chunk.strip()) < 100:
                        continue
                    prompt = [
                        {"role": "user",
                         "content": f"请总结以下文档第{idx}章节的核心内容，重点关注金融SOA相关的规则、结构、表格数据和图片信息：{chunk[:2000]}"}
                    ]
                    summaries.append(f"第{idx}章节摘要：{llm(prompt)}")
                result = "\n\n".join(summaries)
            else:
                import re
                chapter_pattern = re.compile(r'(\d+\.\s+[^\n]+|一、\s+[^\n]+|二、\s+[^\n]+|###\s+[^\n]+)')
                chapters = chapter_pattern.findall(content)
                chapter_str = "文档包含章节：\n" + "\n".join(chapters) if chapters else "文档未识别到明确章节"
                key_content = content[:3000] + "\n[文档中间部分省略]\n" + content[-2000:]
                prompt = [
                    {"role": "user",
                     "content": f"以下是超长金融SOA文档的关键信息（含章节列表、核心片段、表格和图片OCR内容），请总结：1.核心规则要求 2.必填模块 3.表格中的关键数据 4.图片中的有效信息 5.文档结构：\n{chapter_str}\n\n核心片段：{key_content}"}
                ]
                result = llm(prompt)

            self._summaries_cache[cache_key] = result
            return result
        except Exception as e:
            err_msg = f"生成摘要时出错: {str(e)}"
            print(err_msg)
            self.process_errors.append(err_msg)
            return content

    def get_rules_summary(self):
        """生成规则文档的摘要——原有方法，无修改"""
        if not self.rules:
            return "未检测到有效规则文档，将使用默认SOA规则（包含客户背景、建议内容、建议依据、风险提示、费用说明5大模块）"

        rule_summaries = []
        for idx, rule_content in enumerate(self.rules, 1):
            summary = self._summarize_content(rule_content)
            rule_summaries.append(f"### 规则文档{idx}摘要\n{summary}")

        return f"## 行业规则总览\n以下是从{len(self.rules)}个规则文档中提取的核心要求（生成SOA需100%遵守）：\n\n" + "\n\n".join(
            rule_summaries)

    def get_template_structures(self):
        """提取模板文档的结构——原有方法，无修改"""
        if not self.templates:
            return "未检测到有效模板文档，默认SOA结构参考：\n1. 客户背景（姓名、年龄、风险承受能力）\n2. 投资建议内容（产品组合、配置比例）\n3. 建议依据（历史业绩、客户适配性）\n4. 风险提示（市场/产品/流动性风险）\n5. 费用说明（申购费、管理费）"

        template_structures = []
        llm = self._get_llm_instance()

        for idx, template_content in enumerate(self.templates, 1):
            cache_key = hash(template_content[:1000]) + idx
            if cache_key in self._structures_cache:
                template_structures.append(f"### 模板文档{idx}结构\n{self._structures_cache[cache_key]}")
                continue

            try:
                prompt = [
                    {
                        "role": "user",
                        "content": f"请分析以下金融SOA模板文档的结构，输出要求：1. 按「一级章节→二级子模块」格式列出 2. 标注每个模块是否为必填 3. 说明模块间的逻辑顺序 4. 重点标注表格和图片的位置及作用：\n{template_content[:1500]}..."
                    }
                ]
                structure = llm(prompt)
                self._structures_cache[cache_key] = structure
                template_structures.append(f"### 模板文档{idx}结构\n{structure}")
            except Exception as e:
                err_msg = f"提取模板文档{idx}结构时出错: {str(e)}"
                print(err_msg)
                self.process_errors.append(err_msg)
                fallback = f"模板文档{idx}结构（部分提取）：\n{template_content[:800]}..."
                self._structures_cache[cache_key] = fallback
                template_structures.append(fallback)

        return f"## SOA模板结构总览\n以下是从{len(self.templates)}个模板文档中提取的结构框架（生成SOA需严格对齐）：\n\n" + "\n\n".join(
            template_structures)

    def process_all_docs(self):
        """处理文件夹中的所有文档——原有方法，无修改（接口完全不变）"""
        self.rules = []
        self.templates = []
        self.processed_rule_files = []
        self.processed_template_files = []
        self.process_errors = []

        if not os.path.exists(self.docs_folder):
            err_msg = f"文件夹不存在: {self.docs_folder}"
            print(err_msg)
            self.process_errors.append(err_msg)
            return

        supported_extensions = ('.pdf', '.docx')
        for file_name in os.listdir(self.docs_folder):
            file_path = os.path.join(self.docs_folder, file_name)
            if os.path.isdir(file_path):
                print(f"跳过子文件夹: {file_name}")
                continue
            if not file_name.lower().endswith(supported_extensions):
                err_msg = f"跳过不支持的文件格式: {file_name}（仅支持PDF/DOCX）"
                print(err_msg)
                self.process_errors.append(err_msg)
                continue

            content = ""
            if file_name.lower().endswith('.pdf'):
                content = self._process_pdf(file_path)  # 仍调用原有方法名，内部已替换为高级处理
            elif file_name.lower().endswith('.docx'):
                content = self._process_docx(file_path)  # 同上

            if not content.strip():
                err_msg = f"文件无有效文本: {file_name}（可能是扫描件或空白文档）"
                print(err_msg)
                self.process_errors.append(err_msg)
                continue

            if self._is_rule_document(file_name, content):
                self.rules.append(content)
                self.processed_rule_files.append(file_name)
            else:
                self.templates.append(content)
                self.processed_template_files.append(file_name)

            print(f"已处理文件: {file_name}")

        print(
            f"\n处理完成 - 规则文档：{len(self.processed_rule_files)}个，模板文档：{len(self.processed_template_files)}个，错误：{len(self.process_errors)}个")