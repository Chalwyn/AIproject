import streamlit as st
import openai
from getllm import get_llm
import os
from openai import OpenAI
import PyPDF2
from docx import Document
from dotenv import load_dotenv
import os

# 加载 .env 文件中的环境变量
load_dotenv()

client = OpenAI(api_key='OPENAI_API_KEY')

st.set_page_config(layout="wide")


with st.sidebar:
    language = st.selectbox("Choose your language / 选择语言", ["English", "中文"])


#读取word形式的文件函数

from docx import Document

if(language == "中文"):

    # 用官方的api库来完成大模型的调用
    def get_ai_response(prompt):
        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "你是一个专业理财顾问。"},  # system为系统设定
                    {"role": "assistant", "content": file_content},  # 为传入的读取的信息
                    {"role": "user", "content": prompt},  # 输入的用户需求，即问题
                    {"role": "assistant", "content": text1},  # 为传入的读取的信息
                    {"role": "assistant", "content": text2}  # 为传入的读取的信息
                ],
                max_tokens=4000
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"错误: {str(e)}"

    #定义一个将文件翻译成中文的函数

    # Function to translate English to Chinese
    def translate_to_chinese(text):
        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional translator."},
                    {"role": "user", "content": f"Translate the following text to Simplified Chinese:\n\n{text}"}
                ],
                max_tokens=4000
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error: {str(e)}"



    def read_word(file_path):
        try:
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            return f"错误: {str(e)}"

    #输出一下试试

    st.write(translate_to_chinese(read_word("data/ROA S102440(S) Sidney Keith Poole 21-07-2025.docx")))
    text1 = read_word("data/ROA S102440(S) Sidney Keith Poole 21-07-2025.docx")


    #读取pdf格式的文件

    def read_pdf(file_path):
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in range(len(reader.pages)):
                    text += reader.pages[page].extract_text()
            return text
        except Exception as e:
            return f"错误: {str(e)}"


    #输出一下试试

    st.write(translate_to_chinese(read_pdf("data/FSG York Wealth Management 28 January 2025 (V4.2).pdf")))

    text2 = read_pdf("data/FSG York Wealth Management 28 January 2025 (V4.2).pdf")
    #打开data里面的txt格式的文件并保存到file_content
    with open("data/example.txt", "r", encoding="utf-8") as f:
        file_content = f.read()

    st.write(translate_to_chinese(file_content))




    st.page_link("main.py", label="返回主页面")


    st.title("robot")


    prompt = st.chat_input("Say something" , width=300)
    if prompt:
        st.write(prompt)



    if prompt:
        response = get_ai_response(prompt)
        st.write(f"机器人助手: {response}")
    else:
        st.write("请输入问题")

if(language == "English"):

    # Use the official API library to call the large model
    def get_ai_response(prompt):
        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional financial advisor."},  # System setting
                    {"role": "assistant", "content": file_content},  # The information read from the files
                    {"role": "user", "content": prompt},  # User's question
                    {"role": "assistant", "content": text1},  # The Word file content
                    {"role": "assistant", "content": text2}  # The PDF file content
                ],
                max_tokens=4000
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error: {str(e)}"



    #设置一个翻译函数
    # Function to translate Chinese to English
    def translate_to_english(text):
        try:
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a professional translator."},
                    {"role": "user", "content": f"Translate the following text to English:\n\n{text}"}
                ],
                max_tokens=4000
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error: {str(e)}"


    # Read Word file function
    def read_word(file_path):
        try:
            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            return f"Error: {str(e)}"


    # Output the Word content for testing
    st.write(translate_to_english(read_word("data/ROA S102440(S) Sidney Keith Poole 21-07-2025.docx")))
    text1 = read_word("data/ROA S102440(S) Sidney Keith Poole 21-07-2025.docx")


    # Read PDF file function
    def read_pdf(file_path):
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in range(len(reader.pages)):
                    text += reader.pages[page].extract_text()
            return text
        except Exception as e:
            return f"Error: {str(e)}"


    # Output the PDF content for testing
    st.write(translate_to_english(read_pdf("data/FSG York Wealth Management 28 January 2025 (V4.2).pdf")))
    text2 = read_pdf("data/FSG York Wealth Management 28 January 2025 (V4.2).pdf")

    # Open the txt file and save the content
    with open("data/example.txt", "r", encoding="utf-8") as f:
        file_content = f.read()

    st.write(translate_to_english(file_content))






    st.page_link("main.py", label="Back to Main Page")

    st.title("Robot")

    prompt = st.chat_input("Say something", width=300)
    if prompt:
        st.write(prompt)

    if prompt:
        response = get_ai_response(prompt)
        st.write(f"Robot Assistant: {response}")
    else:
        st.write("Please enter your question.")