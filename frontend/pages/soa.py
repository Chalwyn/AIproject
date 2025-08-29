import streamlit as st
import os
import time
from typing import List  # 类型提示增强
import openai
import requests
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter
import os

#get_llm文件内容
import os
import openai
import requests
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter
import os
from openai import OpenAI

client = OpenAI(api_key='OPENAI_API_KEY')



# 全局变量跟踪代理是否已配置
_proxy_configured = False
_session = None


# 配置OpenAI连接，同时保留已设置的API密钥
def configure_openai_proxy():
    global _proxy_configured, _session
    try:
        # 获取当前已设置的API密钥（如果有）
        current_api_key = openai.api_key

        # 只配置一次代理
        if not _proxy_configured:
            # 使用已配置的代理
            proxies = {
                "http": "http://127.0.0.1:7897",
                "https": "http://127.0.0.1:7897"  # 注意这里也是http，因为代理服务器本身是http协议
            }

            _session = requests.Session()
            _session.proxies = proxies
            retry = Retry(total=3, backoff_factor=0.3, status_forcelist=[500, 502, 503, 504])
            adapter = HTTPAdapter(max_retries=retry)
            _session.mount('http://', adapter)
            _session.mount('https://', adapter)

            openai.requestssession = _session
            print(f"已配置代理: {proxies}")
            _proxy_configured = True

        # 如果之前有设置API密钥，重新设置它
        if current_api_key and openai.api_key != current_api_key:
            openai.api_key = current_api_key
    except Exception as e:
        print(f"配置OpenAI连接时出错: {e}")


# 直接从环境文件中读取API密钥
def load_api_key():
    # 尝试从环境变量中获取
    api_key = 'OPENAI_API_KEY'
    # 如果环境变量中没有，尝试从文件中读取
    if not api_key:
        try:
            with open(os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'AiApi.env'), 'r',
                      encoding='utf-8') as f:
                for line in f:
                    if line.strip() and not line.strip().startswith('#'):
                        key, value = line.strip().split('=', 1)
                        if key == 'OPENAI_API_KEY':
                            api_key = value.strip().strip('"').strip("'")
                            break
        except Exception as e:
            print(f"无法从AiApi.env文件加载API密钥: {e}")

    return api_key if api_key else "您的GPT_API_KEY"  # 提供默认值，避免运行时错误


import openai

def get_llm(model_name="gpt-3.5-turbo"):
    # 配置代理
    configure_openai_proxy()

    # 设置API密钥（如果尚未设置）
    if not openai.api_key:
        openai.api_key = load_api_key()

    def summarize_conversation(messages):
        # 添加系统提示词，明确要求总结的格式和重点
        system_prompt = """
        你是一名专业的金融顾问助理，负责总结客户与顾问的对话。请根据以下要求生成总结：
        1. 总结必须包含客户的主要问题或需求
        2. 总结必须包含顾问的核心建议或回复
        3. 保持语言简洁、准确，使用专业金融术语
        4. 不要添加原文中没有的信息
        5. 总结内容的模板大概是：在对话中，客户表达了:...顾问则表达了:...
        """

        # 将系统提示与用户消息合并为一个完整的prompt
        enhanced_messages = [
            {"role": "system", "content": system_prompt}
        ] + messages

        try:
            # 调用chat模型的API接口
            response = openai.chat.completions.create(
                model=model_name,  # 使用chat模型
                messages=enhanced_messages,  # 使用messages作为输入
                temperature=0.3,  # 降低temperature，使输出更稳定
                max_tokens=150  # 限制输出长度
            )

            # 返回生成的文本内容
            return response['choices'][0]['message']['content'].strip()  # 使用 'message' 来获取生成的内容

        except Exception as e:
            return f"错误: {str(e)}"

    return summarize_conversation




import openai

def get_soa_generator(model_name="gpt-3.5-turbo"):
    # 配置代理
    configure_openai_proxy()

    # 设置API密钥（如果尚未设置）
    if not openai.api_key:
        openai.api_key = load_api_key()

    def generate_soa_template(advisor_style, reference_examples, rule_summary=None, template_structure=None):
        """
        生成符合规范的SOA模板Prompt，支持注入规则摘要和模板结构（解决内容简略、偏离模板问题）
        """

        # 构建“规则补充说明”（若有则注入，无则用示例默认规则）
        rule_supplement = ""
        if rule_summary:
            rule_supplement = f"""
            【补充行业规则细节】
            以下是从专业文档中提取的核心规则，需100%遵守，不得遗漏：
            {rule_summary}
            """

        # 构建“模板结构强制要求”（若有则注入，无则明确示例结构优先级）
        structure_requirement = ""
        if template_structure:
            structure_requirement = f"""
            【模板结构强制对齐】
            请严格按照以下提取的模板结构生成，章节标题、顺序、子模块需完全匹配，不得自行删减或调整：
            {template_structure}
            """
        else:
            structure_requirement = """
            【模板结构强制对齐】
            请从参考示例中提取完整模板结构（包括章节层级、子标题、内容模块），生成时需：
            1. 章节数量不少于5个（对应基础规则的5大模块）；
            2. 每个一级章节下需包含至少2个二级子模块（如“客户背景”下含“基本信息”“风险画像”）；
            3. 章节标题与示例保持一致（如示例用“投资建议依据”，不可改为“建议原因”）。
            """

        # 结构化Prompt
        prompt = [
            {
                "role": "system",
                "content": """你是资深金融SOA撰写专家，精通金融行业合规要求，能严格依据规则、模板、顾问风格生成完整文档。
                核心原则：
                1. 内容不简略：每个模块文字量不少于3行，关键模块（如建议依据、风险提示）需包含具体逻辑/数据维度（如“近3年年化收益XX%”“最大回撤XX%”）；
                2. 结构不偏离：完全遵循提供的模板结构，章节顺序、标题、子模块需一一对应；
                3. 合规不遗漏：风险提示必须包含指定语句，客户信息必须用{{占位符}}脱敏。
                若生成内容不符合以上原则，需自动检查并补充完整，无需用户提醒。"""
            },
            {
                "role": "user",
                "content": f"""任务：生成金融顾问专属SOA模板，需同时满足基础规则、行业规则、结构要求、风格约束四大维度，具体要求如下：

            一、基础规则（底线要求，违反则无效）
            1. 必含5大核心模块，缺一不可：
               - 客户背景：含{{客户姓名}}、{{客户年龄}}、{{风险承受能力等级}}、{{投资目标}}（新增，补充基础信息完整性）；
               - 建议内容：含{{产品组合清单}}（至少3类产品）、{{配置比例}}（精确到百分比，如“股票型基金35%”）、{{投资周期建议}}；
               - 建议依据：含{{产品历史业绩}}（近1-3年关键数据）、{{客户目标适配性分析}}（如“匹配客户5年退休规划”）、{{市场环境参考}}；
               - 风险提示：分“市场风险”“产品风险”“流动性风险”3类，每类需举例说明（如“市场风险：A股波动可能导致短期回撤”），且必须包含语句：“本建议非保证收益，过往业绩不代表未来表现”；
               - 费用说明：含{{申购费计算方式}}（如“100万以下1.2%，100万以上0.8%”）、{{管理费标准}}、{{其他费用提示}}（如赎回费、托管费）。
            2. 脱敏要求：所有客户信息、产品具体名称、费用金额用{{占位符}}表示，占位符命名需清晰（如{{客户资产规模}}，不可用{{XXX}}）。

            二、行业规则与模板结构（强制对齐，不得自定义）
            {rule_supplement}
            {structure_requirement}

            三、顾问风格约束（贯穿全文，保持一致性）
            请完全模仿以下顾问写作风格，包括语气、句式、专业术语使用习惯：
            {advisor_style}
            示例：若风格为“严谨合规型”，需多用“根据《XX监管规定》”“经风险评估确认”等表述；若为“通俗易懂型”，需避免复杂术语，用“简单来说”“举个例子”等引导。

            四、参考示例与输出要求
            1. 参考示例：
               {reference_examples}
               （示例仅作参考，若与规则/结构要求冲突，以规则/结构要求为准）
            2. 输出要求：
               - 格式：用Markdown分级标题（# 一级标题，## 二级标题），段落清晰，无杂乱排版；
               - 长度：完整模板文字量尽量详细且多
               - 检查：生成后需自动核对“5大模块是否齐全”“风险提示语句是否包含”“占位符是否规范”，缺失则补充。

            请直接生成完整SOA英文模板，无需额外解释或开场白。"""
            }
        ]

        try:
            # 调用新的chat API接口
            response = client.chat.completions.create(
                model=model_name,  # 使用chat模型
                messages=prompt,  # 使用messages作为输入
                temperature=0.7,  # 温度调节
                max_tokens=1000  # 限制输出长度
            )

            # 返回生成的文本内容
            return response['choices'][0]['message']['content'].strip()

        except Exception as e:
            return f"错误: {str(e)}"

    return generate_soa_template

# 配置OpenAI连接，同时保留已设置的API密钥
def configure_openai_proxy():
    global _proxy_configured, _session
    try:
        # 获取当前已设置的API密钥（如果有）
        current_api_key = openai.api_key

        # 只配置一次代理
        if not _proxy_configured:
            # 使用已配置的代理
            proxies = {
                "http": "http://127.0.0.1:7897",
                "https": "http://127.0.0.1:7897"  # 注意这里也是http，因为代理服务器本身是http协议
            }

            _session = requests.Session()
            _session.proxies = proxies
            retry = Retry(total=3, backoff_factor=0.3, status_forcelist=[500, 502, 503, 504])
            adapter = HTTPAdapter(max_retries=retry)
            _session.mount('http://', adapter)
            _session.mount('https://', adapter)

            openai.requestssession = _session
            print(f"已配置代理: {proxies}")
            _proxy_configured = True

        # 如果之前有设置API密钥，重新设置它
        if current_api_key and openai.api_key != current_api_key:
            openai.api_key = current_api_key
    except Exception as e:
        print(f"配置OpenAI连接时出错: {e}")


#document文件内容

import PyPDF2
from docx import Document
import os
import re
import tempfile
import json
from PIL import Image
import pytesseract
import pdfplumber
import pandas as pd


class DocumentProcessor:
    def __init__(self, docs_folder):
        self.docs_folder = docs_folder
        self.rules = []
        self.templates = []
        # 缓存LLM实例
        self._llm = None
        # 缓存已处理的内容
        self._summaries_cache = {}
        self._structures_cache = {}
        # 新增：记录处理的文件列表与错误信息（用于Streamlit反馈）
        self.processed_rule_files = []  # 处理成功的规则文档名
        self.processed_template_files = []  # 处理成功的模板文档名
        self.process_errors = []  # 解析错误的文档信息（格式："文件名：错误原因"）

        # -------------------------- 新增：图片OCR配置（需用户手动确认路径）--------------------------
        # 1. Windows用户：需将路径改为你的Tesseract安装路径（默认如下，若修改过安装位置需调整）
        self.tesseract_path = r'D:\Program Files\Tesseract-OCR\tesseract.exe'
        # 2. Mac用户：注释上面一行，启用下面一行（默认路径，若用brew安装）
        # self.tesseract_path = '/usr/local/bin/tesseract'

        # 配置Tesseract路径
        try:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_path
            # 验证Tesseract是否可用（避免后续OCR时才报错）
            pytesseract.get_tesseract_version()
        except Exception as e:
            self.process_errors.append(
                f"Tesseract配置错误：{str(e)}，图片OCR功能将不可用。请检查安装路径或重新安装Tesseract（附安装指南：https://github.com/UB-Mannheim/tesseract/wiki）")

    def _get_llm_instance(self):
        """获取LLM实例（单例模式）——原有方法，无修改"""
        if self._llm is None:
            self._llm = get_llm()
        return self._llm

    # -------------------------- 新增：核心工具方法（图片OCR+表格转换）--------------------------
    def _ocr_single_image(self, image_path):
        """单张图片OCR识别，返回结构化文本（标注图片来源）"""
        try:
            # 图片预处理：转为灰度图提升识别率（减少彩色干扰）
            with Image.open(image_path) as img:
                img_gray = img.convert('L')
                # 执行OCR（支持中英双语，避免金融术语识别错误）
                ocr_text = pytesseract.image_to_string(
                    img_gray,
                    lang='chi_sim+eng',  # 中文+英文识别（必须安装对应语言包）
                    config='--psm 6'  # 按段落识别，适合文档类图片
                )
            # 结构化输出：标注图片OCR结果，方便后续LLM识别
            return f"【图片OCR内容】\n{ocr_text.strip()}\n【图片OCR结束】" if ocr_text.strip() else "【图片OCR内容】未识别到有效文字\n【图片OCR结束】"
        except Exception as e:
            return f"【图片OCR内容】识别失败：{str(e)}\n【图片OCR结束】"

    def _table_to_markdown(self, table_data, table_idx, file_type):
        """将表格数据（列表/元组）转为Markdown格式（LLM易识别结构）"""
        try:
            # 处理空表格
            if not table_data or len(table_data) < 1:
                return f"【{file_type}表格{table_idx}】空表格，无数据"

            # 用Pandas整理表格（自动处理空值、对齐列数）
            df = pd.DataFrame(table_data)
            # 若第一行是表头，设为列名（避免表头和内容混为一谈）
            if len(df.columns) > 1 and all(isinstance(cell, str) and len(cell.strip()) > 0 for cell in df.iloc[0]):
                df.columns = df.iloc[0]
                df = df.drop(0).reset_index(drop=True)

            # 空值替换为空白字符串（避免显示NaN）
            df = df.fillna("")
            # 转为Markdown表格
            markdown_table = df.to_markdown(index=False, tablefmt="pipe")
            return f"【{file_type}表格{table_idx}】\n{markdown_table}\n"
        except Exception as e:
            return f"【{file_type}表格{table_idx}】转换失败：{str(e)}\n原始表格数据：{str(table_data[:3])}..."  # 只显示前3行避免过长

    # -------------------------- 新增：PDF图片+表格处理（替换原有纯文本提取）--------------------------
    def _process_pdf_with_table_image(self, file_path):
        """处理PDF中的文本+表格+图片，返回整合后结构化内容"""
        file_name = os.path.basename(file_path)
        content_parts = [f"【PDF文件：{file_name} 内容开始】"]

        try:
            # 1. 提取PDF纯文本（保留基础内容，用于补充表格/图片外的信息）
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_content = ""
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_content += f"【第{page_num}页文本】\n{page_text.strip()}\n"
            if text_content:
                content_parts.append(text_content)

            # 2. 用pdfplumber提取PDF表格（比PyPDF2更精准，保留行列结构）
            with pdfplumber.open(file_path) as pdf:
                table_count = 0
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables()  # 提取当前页所有表格
                    if tables:
                        for table in tables:
                            table_count += 1
                            # 表格转Markdown并添加到内容
                            table_markdown = self._table_to_markdown(table, table_count, "PDF")
                            content_parts.append(f"【第{page_num}页表格】\n{table_markdown}")
                if table_count == 0:
                    content_parts.append("【PDF表格信息】未检测到表格\n")

            # 3. 提取PDF中的图片并执行OCR
            with pdfplumber.open(file_path) as pdf:
                img_count = 0
                for page_num, page in enumerate(pdf.pages, 1):
                    images = page.images  # 提取当前页所有图片
                    if images:
                        for img_info in images:
                            img_count += 1
                            # 提取图片二进制数据，保存为临时文件（OCR完成后删除）
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                                temp_img.write(img_info['stream'].read())
                                temp_img_path = temp_img.name

                            # 执行OCR并添加到内容
                            ocr_result = self._ocr_single_image(temp_img_path)
                            content_parts.append(f"【第{page_num}页图片{img_count}】\n{ocr_result}\n")

                            # 删除临时图片（避免占用磁盘空间）
                            os.unlink(temp_img_path)
                if img_count == 0:
                    content_parts.append("【PDF图片信息】未检测到图片\n")

            # 整合所有内容，返回结构化文本
            content_parts.append(f"【PDF文件：{file_name} 内容结束】")
            return "\n".join(content_parts)

        except Exception as e:
            # 若处理失败，降级为原有纯文本提取（避免完全无法使用）
            err_msg = f"PDF高级处理（表格+图片）失败：{str(e)}，已降级为纯文本提取"
            self.process_errors.append(f"{file_name}：{err_msg}")
            # 调用原有纯文本提取逻辑
            pure_text = self._process_pdf_fallback(file_path)
            return f"【PDF文件：{file_name} 内容（降级纯文本）】\n{pure_text}\n【处理说明】{err_msg}"

    def _process_pdf_fallback(self, file_path):
        """PDF高级处理失败时的降级方案——原有纯文本提取逻辑，无修改"""
        content = ""
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    extracted_text = page.extract_text() or ""
                    content += extracted_text
        except Exception as e:
            err_msg = f"纯文本提取也失败：{str(e)}"
            self.process_errors.append(f"{os.path.basename(file_path)}：{err_msg}")
            content = err_msg
        return content

    # -------------------------- 新增：Word图片+表格处理（替换原有纯文本提取）--------------------------
    def _process_docx_with_table_image(self, file_path):
        """处理Word中的文本+表格+图片，返回整合后结构化内容"""
        file_name = os.path.basename(file_path)
        content_parts = [f"【Word文件：{file_name} 内容开始】"]

        try:
            doc = Document(file_path)

            # 1. 提取Word段落纯文本
            text_content = ""
            for para_idx, para in enumerate(doc.paragraphs, 1):
                para_text = para.text.strip()
                if para_text:
                    text_content += f"【段落{para_idx}】\n{para_text}\n"
            if text_content:
                content_parts.append(text_content)
            else:
                content_parts.append("【Word文本信息】未检测到有效段落文本\n")

            # 2. 提取Word表格
            if doc.tables:
                for table_idx, table in enumerate(doc.tables, 1):
                    # 读取表格数据（行→列）
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]  # 清理换行符
                        table_data.append(row_data)
                    # 表格转Markdown
                    table_markdown = self._table_to_markdown(table_data, table_idx, "Word")
                    content_parts.append(f"【表格{table_idx}】\n{table_markdown}\n")
            else:
                content_parts.append("【Word表格信息】未检测到表格\n")

            # 3. 提取Word中的图片（需通过文档关系获取）
            img_count = 0
            # 遍历文档中的所有关系（图片存储在关系中）
            for rel_id, rel in doc.part.related_parts.items():
                # 判断是否为图片类型（MIME类型匹配）
                if rel.content_type.startswith('image/'):
                    img_count += 1
                    # 保存图片到临时文件
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                        temp_img.write(rel.blob)  # 写入图片二进制数据
                        temp_img_path = temp_img.name

                    # 执行OCR
                    ocr_result = self._ocr_single_image(temp_img_path)
                    content_parts.append(f"【图片{img_count}】\n{ocr_result}\n")

                    # 删除临时图片
                    os.unlink(temp_img_path)
            if img_count == 0:
                content_parts.append("【Word图片信息】未检测到图片\n")

            # 整合所有内容
            content_parts.append(f"【Word文件：{file_name} 内容结束】")
            return "\n".join(content_parts)

        except Exception as e:
            # 降级为原有纯文本提取
            err_msg = f"Word高级处理（表格+图片）失败：{str(e)}，已降级为纯文本提取"
            self.process_errors.append(f"{file_name}：{err_msg}")
            pure_text = self._process_docx_fallback(file_path)
            return f"【Word文件：{file_name} 内容（降级纯文本）】\n{pure_text}\n【处理说明】{err_msg}"

    def _process_docx_fallback(self, file_path):
        """Word高级处理失败时的降级方案——原有纯文本提取逻辑，无修改"""
        content = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                content += para.text + "\n"
        except Exception as e:
            err_msg = f"纯文本提取也失败：{str(e)}"
            self.process_errors.append(f"{os.path.basename(file_path)}：{err_msg}")
            content = err_msg
        return content

    # -------------------------- 原有方法：仅修改内部调用（接口不变，其他代码无需改）--------------------------
    def _process_pdf(self, file_path):
        """原有PDF处理方法——仅修改内部调用，接口（输入输出）完全不变"""
        return self._process_pdf_with_table_image(file_path)  # 调用新增的高级处理方法

    def _process_docx(self, file_path):
        """原有Word处理方法——仅修改内部调用，接口（输入输出）完全不变"""
        return self._process_docx_with_table_image(file_path)  # 调用新增的高级处理方法

    def _is_rule_document(self, file_name, content):
        """判断文档是否为规则文档——原有方法，无修改"""
        rule_keywords = ['规则', '规范', '指引', '要求', '监管', '条例', 'regulation', 'rule', 'guideline',
                         'requirement']
        file_name_lower = file_name.lower()
        content_lower = content.lower() if content else ""

        file_keyword_count = sum(1 for kw in rule_keywords if kw.lower() in file_name_lower)
        if file_keyword_count >= 2:
            return True
        elif file_keyword_count == 1:
            content_keyword_count = sum(1 for kw in rule_keywords if kw.lower() in content_lower)
            return content_keyword_count >= 1
        else:
            content_keyword_count = sum(1 for kw in rule_keywords if kw.lower() in content_lower)
            return content_keyword_count >= 2

    def _summarize_content(self, content):
        """使用LLM生成文档摘要——原有方法，无修改（LLM会自动识别新增的表格/图片标注）"""
        cache_key = hash(content)
        if cache_key in self._summaries_cache:
            return self._summaries_cache[cache_key]

        try:
            llm = self._get_llm_instance()
            if len(content) <= 2000:
                prompt = [
                    {"role": "user",
                     "content": f"请总结以下金融SOA相关文档的核心内容，重点提取：1.规则要求 2.模板结构 3.必填模块 4.表格中的关键数据（如费用标准、风险等级）5.图片OCR中的有效信息，摘要需简洁完整：{content}"}
                ]
                result = llm(prompt)
            elif 2001 <= len(content) <= 10000:
                import re
                chapter_pattern = re.compile(r'(\d+\.\s+|一、\s+|二、\s+|###\s+)')
                chunks = chapter_pattern.split(content)
                full_chunks = []
                current_chunk = ""
                for part in chunks:
                    if chapter_pattern.match(part):
                        if current_chunk:
                            full_chunks.append(current_chunk)
                        current_chunk = part
                    else:
                        current_chunk += part
                if current_chunk:
                    full_chunks.append(current_chunk)
                summaries = []
                for idx, chunk in enumerate(full_chunks, 1):
                    if len(chunk.strip()) < 100:
                        continue
                    prompt = [
                        {"role": "user",
                         "content": f"请总结以下文档第{idx}章节的核心内容，重点关注金融SOA相关的规则、结构、表格数据和图片信息：{chunk[:2000]}"}
                    ]
                    summaries.append(f"第{idx}章节摘要：{llm(prompt)}")
                result = "\n\n".join(summaries)
            else:
                import re
                chapter_pattern = re.compile(r'(\d+\.\s+[^\n]+|一、\s+[^\n]+|二、\s+[^\n]+|###\s+[^\n]+)')
                chapters = chapter_pattern.findall(content)
                chapter_str = "文档包含章节：\n" + "\n".join(chapters) if chapters else "文档未识别到明确章节"
                key_content = content[:3000] + "\n[文档中间部分省略]\n" + content[-2000:]
                prompt = [
                    {"role": "user",
                     "content": f"以下是超长金融SOA文档的关键信息（含章节列表、核心片段、表格和图片OCR内容），请总结：1.核心规则要求 2.必填模块 3.表格中的关键数据 4.图片中的有效信息 5.文档结构：\n{chapter_str}\n\n核心片段：{key_content}"}
                ]
                result = llm(prompt)

            self._summaries_cache[cache_key] = result
            return result
        except Exception as e:
            err_msg = f"生成摘要时出错: {str(e)}"
            print(err_msg)
            self.process_errors.append(err_msg)
            return content

    def get_rules_summary(self):
        """生成规则文档的摘要——原有方法，无修改"""
        if not self.rules:
            return "未检测到有效规则文档，将使用默认SOA规则（包含客户背景、建议内容、建议依据、风险提示、费用说明5大模块）"

        rule_summaries = []
        for idx, rule_content in enumerate(self.rules, 1):
            summary = self._summarize_content(rule_content)
            rule_summaries.append(f"### 规则文档{idx}摘要\n{summary}")

        return f"## 行业规则总览\n以下是从{len(self.rules)}个规则文档中提取的核心要求（生成SOA需100%遵守）：\n\n" + "\n\n".join(
            rule_summaries)

    def get_template_structures(self):
        """提取模板文档的结构——原有方法，无修改"""
        if not self.templates:
            return "未检测到有效模板文档，默认SOA结构参考：\n1. 客户背景（姓名、年龄、风险承受能力）\n2. 投资建议内容（产品组合、配置比例）\n3. 建议依据（历史业绩、客户适配性）\n4. 风险提示（市场/产品/流动性风险）\n5. 费用说明（申购费、管理费）"

        template_structures = []
        llm = self._get_llm_instance()

        for idx, template_content in enumerate(self.templates, 1):
            cache_key = hash(template_content[:1000]) + idx
            if cache_key in self._structures_cache:
                template_structures.append(f"### 模板文档{idx}结构\n{self._structures_cache[cache_key]}")
                continue

            try:
                prompt = [
                    {
                        "role": "user",
                        "content": f"请分析以下金融SOA模板文档的结构，输出要求：1. 按「一级章节→二级子模块」格式列出 2. 标注每个模块是否为必填 3. 说明模块间的逻辑顺序 4. 重点标注表格和图片的位置及作用：\n{template_content[:1500]}..."
                    }
                ]
                structure = llm(prompt)
                self._structures_cache[cache_key] = structure
                template_structures.append(f"### 模板文档{idx}结构\n{structure}")
            except Exception as e:
                err_msg = f"提取模板文档{idx}结构时出错: {str(e)}"
                print(err_msg)
                self.process_errors.append(err_msg)
                fallback = f"模板文档{idx}结构（部分提取）：\n{template_content[:800]}..."
                self._structures_cache[cache_key] = fallback
                template_structures.append(fallback)

        return f"## SOA模板结构总览\n以下是从{len(self.templates)}个模板文档中提取的结构框架（生成SOA需严格对齐）：\n\n" + "\n\n".join(
            template_structures)

    def process_all_docs(self):
        """处理文件夹中的所有文档——原有方法，无修改（接口完全不变）"""
        self.rules = []
        self.templates = []
        self.processed_rule_files = []
        self.processed_template_files = []
        self.process_errors = []

        if not os.path.exists(self.docs_folder):
            err_msg = f"文件夹不存在: {self.docs_folder}"
            print(err_msg)
            self.process_errors.append(err_msg)
            return

        supported_extensions = ('.pdf', '.docx')
        for file_name in os.listdir(self.docs_folder):
            file_path = os.path.join(self.docs_folder, file_name)
            if os.path.isdir(file_path):
                print(f"跳过子文件夹: {file_name}")
                continue
            if not file_name.lower().endswith(supported_extensions):
                err_msg = f"跳过不支持的文件格式: {file_name}（仅支持PDF/DOCX）"
                print(err_msg)
                self.process_errors.append(err_msg)
                continue

            content = ""
            if file_name.lower().endswith('.pdf'):
                content = self._process_pdf(file_path)  # 仍调用原有方法名，内部已替换为高级处理
            elif file_name.lower().endswith('.docx'):
                content = self._process_docx(file_path)  # 同上

            if not content.strip():
                err_msg = f"文件无有效文本: {file_name}（可能是扫描件或空白文档）"
                print(err_msg)
                self.process_errors.append(err_msg)
                continue

            if self._is_rule_document(file_name, content):
                self.rules.append(content)
                self.processed_rule_files.append(file_name)
            else:
                self.templates.append(content)
                self.processed_template_files.append(file_name)

            print(f"已处理文件: {file_name}")

        print(
            f"\n处理完成 - 规则文档：{len(self.processed_rule_files)}个，模板文档：{len(self.processed_template_files)}个，错误：{len(self.process_errors)}个")


st.page_link("main.py", label="返回主页面")


st.title("soa")


st.title("📊 个性化投资建议声明书 (SOA) 生成工具")
st.divider()  # 分割线提升视觉体验

# -------------------------- 1. 会话状态初始化（增强缓存与状态管理）--------------------------
# 补充：缓存处理后的规则摘要和模板结构，避免重复计算
if 'session_state_init' not in st.session_state:
    # 基础配置
    st.session_state['advisor_style'] = """
- 风险提示模块必须用"1. 2. 3."分点表述，每个风险点后补充1句具体场景例子（如"市场风险：若A股下跌20%，股票基金可能回撤15%"）
- 建议依据模块需包含"基于{{产品名称}}近{{X}}年历史数据，{{关键指标}}（如年化收益、最大回撤）优于同类产品{{XX}}%"的标准化表述
- 整体语言正式合规，避免"大概""可能"等模糊词汇，段落间用"### 模块名称"明确分隔
- 费用说明需拆分"申购费""管理费""赎回费"三类，每类标注计算方式（如"申购费：100万以下1.2%，100万以上0.8%"）
    """
    st.session_state['reference_examples'] = """
### 风险提示
1. 市场风险：本组合中股票基金占比35%，若市场出现系统性下跌（如2022年沪深300指数下跌21%），可能导致组合净值回撤18%-22%。
2. 产品风险：本次推荐的{{债券基金名称}}虽为中低风险，但仍存在信用风险——若持仓债券发行人（如{{发行人名称}}）违约，可能影响收益兑付。
3. 流动性风险：{{封闭式基金名称}}锁定期为1年，锁定期内无法赎回，需客户匹配长期资金规划。

### 建议依据
基于{{股票基金名称}}近5年历史数据，其年化收益达12.3%，较同类基金平均水平（8.5%）高出3.8个百分点；{{债券基金名称}}近3年最大回撤仅2.1%，符合客户"稳健增值"的投资目标，适配其C3（平衡型）风险承受能力等级。
    """
    # 新增：缓存处理后的规则和结构，避免重复计算
    st.session_state['saved_rules_summary'] = ""
    st.session_state['saved_template_structures'] = ""
    # 新增：LLM参数（用户可调节，控制生成风格）
    st.session_state['llm_temperature'] = 0.3  # 初始值：低温度=更严谨
    st.session_state['llm_max_tokens'] = 2500  # 初始值：确保生成完整模板
    # 新增：文档处理的详细反馈（错误信息、处理文件列表）
    st.session_state['doc_process_errors'] = []
    st.session_state['processed_rule_files'] = []
    st.session_state['processed_template_files'] = []
    # 标记初始化完成
    st.session_state['session_state_init'] = True

# 保留原有核心会话状态
if 'generated_template' not in st.session_state:
    st.session_state['generated_template'] = ""
if 'docs_processor' not in st.session_state:
    st.session_state['docs_processor'] = None
if 'docs_folder' not in st.session_state:
    st.session_state['docs_folder'] = ""


# -------------------------- 2. 用户输入区域优化（增强引导与可配置性）--------------------------
# 2.1 顾问风格设置（增加示例提示）
st.subheader("🎯 顾问风格配置", help="描述顾问的写作习惯、格式要求，越详细越精准")
st.session_state['advisor_style'] = st.text_area(
    label="顾问写作风格描述（示例：风险提示分点、费用说明需含计算方式）",
    value=st.session_state['advisor_style'],
    height=180,
    help="可填写：语言正式度、模块格式要求（如分点/段落）、必含表述（如监管依据引用）"
)

# 2.2 参考示例（增加格式引导）
st.subheader("📑 参考示例输入", help="粘贴顾问以往的SOA片段，确保生成风格一致")
st.session_state['reference_examples'] = st.text_area(
    label="SOA参考片段（建议包含风险提示、建议依据模块）",
    value=st.session_state['reference_examples'],
    height=220,
    help="示例需包含真实模块结构（如### 风险提示），占位符用{{变量名}}表示（如{{客户姓名}}）"
)

# 2.3 文档文件夹处理（增强路径验证与反馈）
st.subheader("📂 SOA规则/模板文档管理", help="上传包含行业规则、SOA模板的PDF/DOCX文件")
col1, col2 = st.columns([3, 1])
with col1:
    st.session_state['docs_folder'] = st.text_input(
        label="文件夹路径（Windows：D:\\soa_docs；Mac/Linux：/Users/soa_docs）",
        value=st.session_state['docs_folder'],
        help="文件夹需包含：规则文档（含'规则''规范'关键词）、模板文档（SOA示例）"
    )
with col2:
    # 新增：快速清空路径按钮
    if st.button("清空路径", use_container_width=True):
        st.session_state['docs_folder'] = ""
        st.session_state['saved_rules_summary'] = ""
        st.session_state['saved_template_structures'] = ""
        st.session_state['doc_process_errors'] = []
        st.session_state['processed_rule_files'] = []
        st.session_state['processed_template_files'] = []
        st.rerun()

# 2.4 LLM生成参数（新增：用户可调节，解决内容简略问题）
st.subheader("⚙️ LLM生成参数", help="调节生成内容的严谨度与长度")
col_temp, col_tokens = st.columns(2)
with col_temp:
    st.session_state['llm_temperature'] = st.slider(
        label="温度（0=严谨固定，1=灵活多样）",
        min_value=0.0,
        max_value=1.0,
        step=0.1,
        value=st.session_state['llm_temperature'],
        help="生成SOA建议设为0.2-0.4，避免偏离规则"
    )
with col_tokens:
    st.session_state['llm_max_tokens'] = st.slider(
        label="最大输出长度（字符数）",
        min_value=1000,
        max_value=4000,
        step=100,
        value=st.session_state['llm_max_tokens'],
        help="建议设为2000-3000，确保完整包含5大模块"
    )


# -------------------------- 3. 文档处理功能增强（详细反馈+错误可视化）--------------------------
if st.button("🔍 处理文件夹中的文档", use_container_width=True):
    # 1. 基础验证：路径非空且存在
    if not st.session_state['docs_folder'].strip():
        st.warning("❌ 请先输入有效的文件夹路径！")
    elif not os.path.exists(st.session_state['docs_folder']):
        st.error(f"❌ 文件夹不存在：{st.session_state['docs_folder']}，请检查路径是否正确")
    elif not os.path.isdir(st.session_state['docs_folder']):
        st.error(f"❌ 输入的不是文件夹路径：{st.session_state['docs_folder']}")
    else:
        try:
            with st.spinner('📄 正在解析文件夹中的文档（PDF/DOCX）...'):
                # 初始化文档处理器，并新增错误收集
                processor =DocumentProcessor(st.session_state['docs_folder'])
                processor.processed_rule_files = []  # 处理成功的规则文档名
                processor.processed_template_files = []  # 处理成功的模板文档名
                processor.process_errors = []  # 解析错误的文档信息
                processor.process_all_docs()

                # 保存处理结果到会话状态
                st.session_state['docs_processor'] = processor
                st.session_state['saved_rules_summary'] = processor.get_rules_summary()
                st.session_state['saved_template_structures'] = processor.get_template_structures()
                st.session_state['doc_process_errors'] = processor.process_errors
                st.session_state['processed_rule_files'] = processor.processed_rule_files
                st.session_state['processed_template_files'] = processor.processed_template_files

                # 2. 显示处理结果（详细反馈）
                st.success(f"✅ 文档处理完成！")
                # 显示处理的文件列表
                with st.expander("查看处理详情", expanded=True):
                    # 规则文档
                    st.markdown(f"**📜 处理成功的规则文档（{len(processor.processed_rule_files)}个）**")
                    if processor.processed_rule_files:
                        for idx, file in enumerate(processor.processed_rule_files, 1):
                            st.markdown(f"{idx}. {file}")
                    else:
                        st.markdown("暂无规则文档（规则文档需含'规则''规范''guideline'等关键词）")

                    # 模板文档
                    st.markdown(f"**📋 处理成功的模板文档（{len(processor.processed_template_files)}个）**")
                    if processor.processed_template_files:
                        for idx, file in enumerate(processor.processed_template_files, 1):
                            st.markdown(f"{idx}. {file}")
                    else:
                        st.markdown("暂无模板文档（非规则文档默认归类为模板文档）")

                    # 错误信息
                    if processor.process_errors:
                        st.markdown(f"**⚠️ 解析失败的文档（{len(processor.process_errors)}个）**")
                        for err in processor.process_errors:
                            st.markdown(f"❌ {err}")

                # 3. 预览提取的规则和结构（让用户确认是否正确）
                with st.expander("预览提取的规则摘要与模板结构", expanded=False):
                    if st.session_state['saved_rules_summary']:
                        st.markdown("**📜 行业规则摘要**")
                        st.markdown(st.session_state['saved_rules_summary'])
                    else:
                        st.markdown("⚠️ 未提取到规则摘要（请确保规则文档包含有效文本）")

                    if st.session_state['saved_template_structures']:
                        st.markdown("**📋 模板结构分析**")
                        st.markdown(st.session_state['saved_template_structures'])
                    else:
                        st.markdown("⚠️ 未提取到模板结构（请确保模板文档包含有效章节）")

        except Exception as e:
            st.error(f"❌ 文档处理失败：{str(e)}（建议检查文件夹权限或文档格式）")


# -------------------------- 4. 生成模板逻辑优化（对接增强型Prompt）--------------------------
if st.button("🚀 生成个性化SOA模板", use_container_width=True, type="primary"):
    # 1. 前置验证
    if not st.session_state['advisor_style'].strip():
        st.warning("❌ 请先填写顾问风格配置！")
    elif not st.session_state['reference_examples'].strip():
        st.warning("❌ 请先填写SOA参考示例！")
    else:
        try:
            with st.spinner('🤖 正在调用LLM生成SOA模板（请耐心等待）...'):
                # 2. 获取LLM实例（传入用户配置的参数）
                soa_generator = get_soa_generator(model_name="gpt-3.5-turbo")

                # 3. 构建增强型Prompt（调用之前优化的generate_soa_template函数）
                st.session_state['generated_template'] = soa_generator(
                    advisor_style=st.session_state['advisor_style'],
                    reference_examples=st.session_state['reference_examples'],
                    rule_summary=st.session_state['saved_rules_summary'],
                    template_structure=st.session_state['saved_template_structures']
                )

                # 4. 调用LLM生成模板（处理流式输出或直接调用，根据get_soa_generator实现调整）
                # # 假设soa_generator接受结构化Prompt（system+user），返回生成文本
                # st.session_state['generated_template'] = soa_generator(enhanced_prompt)

                # 5. 生成成功反馈
                st.success("✅ 个性化SOA模板生成完成！")

        except Exception as e:
            # 细化错误类型（帮助用户排查）
            if "API key" in str(e) or "authentication" in str(e).lower():
                st.error(f"❌ LLM调用失败：API密钥无效或未配置，请检查密钥设置")
            elif "timeout" in str(e).lower():
                st.error(f"❌ LLM调用超时：网络不稳定或LLM响应缓慢，建议重试")
            elif "context length" in str(e).lower():
                st.error(f"❌ 上下文长度超限：请减少参考示例字数或降低max_tokens值")
            else:
                st.error(f"❌ 模板生成失败：{str(e)}")


# -------------------------- 5. 结果显示与下载优化（Markdown渲染+编辑功能）--------------------------
if st.session_state['generated_template']:
    st.subheader("📄 生成的SOA模板", help="可直接复制使用，或编辑后下载")

    # 5.1 用Markdown渲染模板（更直观，支持分级标题）
    st.markdown("### 模板预览（支持Markdown格式）")
    with st.container(border=True):
        st.markdown(st.session_state['generated_template'])

    # 5.2 提供编辑功能（用户可修改后下载）
    st.markdown("### 模板编辑（修改后点击下载）")
    edited_template = st.text_area(
        label="编辑SOA模板",
        value=st.session_state['generated_template'],
        height=300,
        help="可修改占位符、补充模块内容，保存后点击下载"
    )

    # 5.3 下载功能（支持Markdown和TXT格式）
    st.markdown("### 下载模板")
    col_md, col_txt = st.columns(2)
    with col_md:
        st.download_button(
            label="下载为Markdown文件（.md）",
            data=edited_template,
            file_name=f"soa_template_{time.strftime('%Y%m%d%H%M%S')}.md",
            mime="text/markdown",
            use_container_width=True
        )
    with col_txt:
        st.download_button(
            label="下载为文本文件（.txt）",
            data=edited_template,
            file_name=f"soa_template_{time.strftime('%Y%m%d%H%M%S')}.txt",
            mime="text/plain",
            use_container_width=True
        )


# -------------------------- 6. 使用指南优化（补充文档处理和参数说明）--------------------------
with st.expander("📖 使用指南（点击查看详细步骤）", expanded=False):
    st.markdown("""
    ### 完整操作流程
    1. **配置顾问风格**  
       - 填写顾问的写作习惯（如风险提示分点、语言正式度）  
       - 必含格式要求（如"费用说明需拆分申购费/管理费"）

    2. **上传参考示例**  
       - 粘贴顾问以往的SOA片段（至少包含1-2个核心模块，如风险提示、建议依据）  
       - 占位符用`{{变量名}}`表示（如`{{客户姓名}}`、`{{风险承受能力等级}}`）

    3. **处理规则/模板文档（可选但推荐）**  
       - 输入包含行业规则（如"金融SOA监管规范.pdf"）和SOA模板（如"SOA模板示例.docx"）的文件夹路径  
       - 点击【处理文件夹中的文档】，查看提取的规则摘要和模板结构（确保符合预期）

    4. **调节LLM参数**  
       - 温度：建议设为0.2-0.4（越低越严谨，避免偏离规则）  
       - 最大长度：建议设为2000-3000（确保完整包含5大模块：客户背景、建议内容、建议依据、风险提示、费用说明）

    5. **生成与使用模板**  
       - 点击【生成个性化SOA模板】，等待LLM处理  
       - 预览模板后可编辑（如补充占位符、调整模块顺序），最后下载使用

    ### 注意事项
    - 文档格式：仅支持PDF和DOCX，确保文档可提取文本（扫描件需先OCR处理）  
    - 规则文档：文件名或内容需含"规则""规范""guideline"等关键词，否则会被归类为模板文档  
    - LLM调用：确保API密钥配置正确（如OpenAI密钥、本地化LLM服务正常）
    """)
